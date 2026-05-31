from app.parsers.base import BaseParser, ParsedDocument


class DocxParser(BaseParser):
    """Parser for DOCX (Microsoft Word) documents.

    Extracts text content from DOCX files using python-docx,
    preserving paragraph structure as sections.
    """

    async def parse(self, file_path: str) -> ParsedDocument:
        """Parse a DOCX file and extract text content.

        Reads all paragraphs from the document, identifying headings
        by their style and preserving paragraph boundaries.

        Args:
            file_path: Path to the DOCX file on disk.

        Returns:
            A ParsedDocument with full text and heading sections.
        """
        try:
            from docx import Document

            doc = Document(file_path)

            paragraphs: list[str] = []
            sections: list[str] = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                paragraphs.append(text)

                if paragraph.style and "heading" in paragraph.style.name.lower():
                    sections.append(text)

            content = "\n\n".join(paragraphs)

            return ParsedDocument(
                content=content,
                metadata={
                    "file_type": "docx",
                    "paragraph_count": len(paragraphs),
                    "heading_count": len(sections),
                    "char_count": len(content),
                },
                sections=sections,
            )
        except ImportError:
            return ParsedDocument(
                content="",
                metadata={"error": "python-docx not installed", "file_type": "docx"},
                sections=[],
            )
